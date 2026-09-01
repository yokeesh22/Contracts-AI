"""The issues list: a dense tabular summary of the review for circulation.

The companion to the tracked-changes export. That file goes to the vendor; this
one goes to the deal team and to whoever has to approve the risk. Tabular is the
right shape here even though the on-screen findings are cards, because this is
read as a summary rather than worked through clause by clause.
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CLASSIFICATION_LABELS = {
    "UNACCEPTABLE": "Unacceptable",
    "MISSING": "Missing Protection",
    "NEGOTIABLE": "Negotiable",
    "ACCEPTABLE": "Acceptable",
}

HEADER_FILL = {
    "UNACCEPTABLE": "C00000",
    "MISSING": "7030A0",
    "NEGOTIABLE": "ED7D31",
    "ACCEPTABLE": "548235",
}

CELL_FILL = {
    "UNACCEPTABLE": "FFDCE1",
    "MISSING": "EDE1F5",
    "NEGOTIABLE": "FCE4D6",
    "ACCEPTABLE": "E2EFDA",
}

# Contract vocabulary is full of acronyms; .title() renders these as "Ai
# Training Data" and "Incorporation By Url", which reads as a defect.
ACRONYMS = {"AI", "IP", "URL", "DPA", "SLA", "GAI", "PI", "US", "EU"}
LOWERCASE_WORDS = {"by", "of", "on", "to", "for", "and", "the", "in"}


def humanise_clause_type(value: str) -> str:
    words = []
    for i, word in enumerate((value or "").split("_")):
        if not word:
            continue
        if word in ACRONYMS:
            words.append(word)
        elif i > 0 and word.lower() in LOWERCASE_WORDS:
            words.append(word.lower())
        else:
            words.append(word.capitalize())
    return " ".join(words)


STATUS_LABELS = {
    "suggested": "Proposed",
    "accepted": "Accepted",
    "rejected": "Rejected (not in redline)",
    "modified": "Edited by reviewer",
}


def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _write_cell(cell, text: str, size: int = 8, bold: bool = False, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def generate_issues_list_docx(review, version, redlines: list) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    title = doc.add_heading("Contract Review - Issues List", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    counts: dict[str, int] = {}
    for redline in redlines:
        counts[redline.classification] = counts.get(redline.classification, 0) + 1

    info = doc.add_table(rows=6, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Contract", review.name),
        ("Counterparty", review.counterparty or "-"),
        ("Playbook", review.playbook.name if review.playbook else "-"),
        ("Round", f"{version.round_number}"),
        ("Reviewed", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Findings", str(len(redlines))),
    ]
    for i, (label, value) in enumerate(rows):
        _write_cell(info.rows[i].cells[0], label, size=9, bold=True)
        _write_cell(info.rows[i].cells[1], value, size=9)

    doc.add_paragraph()

    doc.add_heading("Summary", level=2)
    summary = doc.add_table(rows=1, cols=len(CLASSIFICATION_LABELS))
    summary.style = "Table Grid"
    for i, (key, label) in enumerate(CLASSIFICATION_LABELS.items()):
        cell = summary.rows[0].cells[i]
        _set_cell_bg(cell, HEADER_FILL[key])
        _write_cell(
            cell,
            f"{label}\n{counts.get(key, 0)}",
            size=9,
            bold=True,
            color=RGBColor(0xFF, 0xFF, 0xFF),
        )

    if version.doc_kind != "docx":
        note = doc.add_paragraph()
        note.add_run(
            "This contract was uploaded as a PDF. The accompanying redline file is "
            "reconstructed from extracted text and does not preserve the original "
            "formatting."
        ).italic = True

    doc.add_paragraph()

    doc.add_heading("Findings", level=2)
    headers = [
        "#",
        "Clause",
        "Type",
        "Assessment",
        "Issue / Rationale",
        "Proposed Change",
        "Status",
    ]
    widths = [Cm(0.8), Cm(2.6), Cm(2.6), Cm(2.2), Cm(5.2), Cm(5.6), Cm(2.4)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "1F4E79")
        _write_cell(cell, header, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.width = width

    for idx, redline in enumerate(redlines, start=1):
        classification = redline.classification
        row = table.add_row()

        clause_label = redline.clause_ref or "-"
        if redline.clause_title:
            clause_label = (
                f"{redline.clause_ref} {redline.clause_title}"
                if redline.clause_ref
                else redline.clause_title
            )

        if classification == "MISSING":
            proposed = f"ADD: {redline.proposed_text or '-'}"
        elif redline.status == "rejected":
            proposed = "No change - suggestion rejected"
        else:
            proposed = redline.proposed_text or "-"

        values = [
            str(idx),
            clause_label,
            humanise_clause_type(redline.clause_type) or "-",
            CLASSIFICATION_LABELS.get(classification, classification),
            redline.rationale or "-",
            proposed,
            STATUS_LABELS.get(redline.status, redline.status),
        ]

        for i, (value, width) in enumerate(zip(values, widths)):
            cell = row.cells[i]
            _write_cell(cell, value)
            cell.width = width
            if i == 3 and classification in CELL_FILL:
                _set_cell_bg(cell, CELL_FILL[classification])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
