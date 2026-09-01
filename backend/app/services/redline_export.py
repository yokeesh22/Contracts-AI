"""Write the redline back out as a Word file with real tracked changes.

Two paths, because the source formats differ in what they can carry:

  * .docx source - revision marks are written into the original document, so
    the vendor receives their own paper back with formatting intact. This is
    the deliverable a legal team actually sends.

  * .pdf source - a PDF has no paragraph structure to write revision marks
    into, so a fresh document is reconstructed from the extracted blocks. The
    edits and comments are correct; the vendor's layout is not preserved. The
    caller is expected to tell the user which of the two they got.

Only redlines with status "suggested", "accepted" or "modified" are exported.
Rejected ones leave the clause exactly as the vendor wrote it.
"""

import copy
import io
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from xml.etree import ElementTree as ET

from .diff_util import diff_ops
from .document_model import load_blocks

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
W = f"{{{W_NS}}}"

ET.register_namespace("w", W_NS)
ET.register_namespace("r", R_NS)

COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
COMMENTS_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)

# Only decisions a human actually made reach the counterparty. A "suggested"
# redline is one nobody has looked at yet, and putting those in the file means
# sending the vendor edits your own team has not agreed to - the export has to
# reflect the review, not the analysis that preceded it.
EXPORTABLE = ("accepted", "modified")


def _now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _el(tag: str, **attrs) -> ET.Element:
    e = ET.Element(f"{W}{tag}")
    for k, v in attrs.items():
        e.set(f"{W}{k}", str(v))
    return e


def _run(text: str, rpr: ET.Element | None, deleted: bool = False) -> ET.Element:
    """A run holding text. Deleted runs use w:delText, which is what makes Word
    render them as struck-through rather than as ordinary text."""
    r = ET.Element(f"{W}r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = ET.SubElement(r, f"{W}delText" if deleted else f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


class _Revision:
    """Allocates the ids Word requires to be unique across the document."""

    def __init__(self, author: str, start_at: int = 1000):
        self.author = author
        self._next = start_at
        self.date = _now()

    def take(self) -> int:
        self._next += 1
        return self._next

    def ins(self) -> ET.Element:
        return _el("ins", id=self.take(), author=self.author, date=self.date)

    def dele(self) -> ET.Element:
        return _el("del", id=self.take(), author=self.author, date=self.date)


def _first_rpr(para: ET.Element) -> ET.Element | None:
    run = para.find(f"{W}r")
    return None if run is None else run.find(f"{W}rPr")


def _apply_diff_to_paragraph(
    para: ET.Element, original: str, proposed: str, rev: _Revision
) -> None:
    """Rewrite one paragraph as interleaved kept / deleted / inserted runs.

    The paragraph's own properties and its first run's formatting are carried
    over, so the redlined clause still looks like the rest of the contract.
    Intra-paragraph formatting changes (a bolded word mid-clause) are not
    preserved - an acceptable trade for a clause being wholly rewritten.
    """
    rpr = _first_rpr(para)
    ppr = para.find(f"{W}pPr")

    for child in list(para):
        para.remove(child)
    if ppr is not None:
        para.append(ppr)

    for op in diff_ops(original, proposed):
        if op["op"] == "equal":
            para.append(_run(op["text"], rpr))
        elif op["op"] == "delete":
            wrapper = rev.dele()
            wrapper.append(_run(op["text"], rpr, deleted=True))
            para.append(wrapper)
        elif op["op"] == "insert":
            wrapper = rev.ins()
            wrapper.append(_run(op["text"], rpr))
            para.append(wrapper)


def _delete_paragraph(para: ET.Element, rev: _Revision) -> None:
    """Mark a whole paragraph deleted, including its paragraph mark, so Word
    removes the empty line when the change is accepted."""
    rpr = _first_rpr(para)
    text = "".join(t.text or "" for t in para.iter(f"{W}t"))
    ppr = para.find(f"{W}pPr")

    for child in list(para):
        para.remove(child)
    if ppr is None:
        ppr = ET.Element(f"{W}pPr")
    para.append(ppr)

    # Marking the paragraph mark itself deleted lives in pPr/rPr/del.
    mark_rpr = ppr.find(f"{W}rPr")
    if mark_rpr is None:
        mark_rpr = ET.SubElement(ppr, f"{W}rPr")
    mark_rpr.append(rev.dele())

    if text.strip():
        wrapper = rev.dele()
        wrapper.append(_run(text, rpr, deleted=True))
        para.append(wrapper)


def _add_comment(
    para: ET.Element, comments_root: ET.Element, cid: int, author: str, text: str
) -> None:
    """Attach a margin comment spanning the paragraph.

    The rationale is the part of a redline that persuades; a bare edit with no
    explanation reads as an ultimatum. Matches how the Scribe sample argued its
    position in comments rather than in the text.
    """
    para.insert(0, _el("commentRangeStart", id=cid))
    para.append(_el("commentRangeEnd", id=cid))

    ref_run = ET.SubElement(para, f"{W}r")
    rpr = ET.SubElement(ref_run, f"{W}rPr")
    ET.SubElement(rpr, f"{W}rStyle").set(f"{W}val", "CommentReference")
    ET.SubElement(ref_run, f"{W}commentReference").set(f"{W}id", str(cid))

    comment = _el("comment", id=cid, author=author, date=_now(), initials="AI")
    cpara = ET.SubElement(comment, f"{W}p")
    crun = ET.SubElement(cpara, f"{W}r")
    ctext = ET.SubElement(crun, f"{W}t")
    ctext.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    ctext.text = text
    comments_root.append(comment)


def _exportable(redlines) -> list:
    out = []
    for r in redlines:
        if r.status not in EXPORTABLE:
            continue
        if r.classification == "ACCEPTABLE" and not r.proposed_text:
            continue
        if not r.proposed_text:
            continue
        out.append(r)
    return out


def exportable_redlines(redlines) -> list:
    """The redlines that will actually be written into the file.

    Public so callers can report the count before a download rather than letting
    someone discover their export was empty by opening it in Word.
    """
    return _exportable(redlines)


def _body_paragraphs(body: ET.Element) -> dict[int, ET.Element]:
    """Map body child position -> paragraph element.

    Keyed on position because that is the same coordinate document_model
    recorded as `w_index` during extraction; matching on text would break the
    moment the model proposed changing that text.
    """
    return {
        i: child
        for i, child in enumerate(body)
        if child.tag == f"{W}p"
    }


def _ensure_comments_part(zin: zipfile.ZipFile) -> tuple[ET.Element, bool]:
    """Return the comments root, and whether the part already existed."""
    if "word/comments.xml" in zin.namelist():
        return ET.fromstring(zin.read("word/comments.xml")), True
    return ET.Element(f"{W}comments"), False


def _register_comments_part(zin: zipfile.ZipFile) -> tuple[bytes, bytes]:
    """Add the content-type override and relationship a new comments part needs.

    Word silently discards comments whose part is not declared in both places,
    which presents as an export that "loses" every rationale.
    """
    ct = ET.fromstring(zin.read("[Content_Types].xml"))
    has_ct = any(
        o.get("PartName") == "/word/comments.xml" for o in ct.findall(f"{{{CT_NS}}}Override")
    )
    if not has_ct:
        override = ET.SubElement(ct, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", COMMENTS_CT)

    rels = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    has_rel = any(r.get("Type") == COMMENTS_REL for r in rels)
    if not has_rel:
        used = {r.get("Id") for r in rels}
        rid = next(f"rId{n}" for n in range(900, 999) if f"rId{n}" not in used)
        rel = ET.SubElement(rels, f"{{{rel_ns}}}Relationship")
        rel.set("Id", rid)
        rel.set("Type", COMMENTS_REL)
        rel.set("Target", "comments.xml")

    return ET.tostring(ct, xml_declaration=True, encoding="UTF-8"), ET.tostring(
        rels, xml_declaration=True, encoding="UTF-8"
    )


def export_tracked_docx(version, redlines, author: str = "Contracts.AI") -> bytes:
    """Write revision marks into the original .docx and return the new file.

    Takes the version rather than the negotiation: from round two the revisions
    have to be written into the file the counterparty last sent us, not into the
    paper they opened with.
    """
    blocks = load_blocks(version.blocks_json)
    by_index = {b["index"]: b for b in blocks}
    rev = _Revision(author)

    with zipfile.ZipFile(version.file_path) as zin:
        doc_root = ET.fromstring(zin.read("word/document.xml"))
        comments_root, comments_existed = _ensure_comments_part(zin)
        ct_bytes, rels_bytes = _register_comments_part(zin)
        other_parts = {
            name: zin.read(name)
            for name in zin.namelist()
            if name
            not in (
                "word/document.xml",
                "word/comments.xml",
                "[Content_Types].xml",
                "word/_rels/document.xml.rels",
            )
        }

    body = doc_root.find(f"{W}body")
    paragraphs = _body_paragraphs(body)

    # Comment ids must not collide with any already in the file.
    existing_ids = [
        int(c.get(f"{W}id", "0")) for c in comments_root.findall(f"{W}comment")
    ]
    next_cid = (max(existing_ids) + 1) if existing_ids else 1

    orphans = []
    # Paragraphs already rewritten. A second diff applied to the same paragraph
    # would be computed against text that is no longer there, and would discard
    # the first edit — losing it from the file the counterparty receives. The
    # engine groups findings per clause to avoid this; a reviewer adding their
    # own redline over an existing one can still collide, so the later edit is
    # appended instead of overwriting.
    rewritten: set[int] = set()

    for redline in _exportable(redlines):
        if redline.block_start is None:
            orphans.append(redline)
            continue

        w_indices = []
        for i in range(redline.block_start, (redline.block_end or redline.block_start) + 1):
            block = by_index.get(i)
            if block and block.get("w_index") is not None:
                if block["w_index"] not in w_indices:
                    w_indices.append(block["w_index"])

        targets = [paragraphs[w] for w in w_indices if w in paragraphs]
        if not targets:
            # The clause sat in a table, or the block model and the document
            # have diverged. Append it rather than dropping it silently.
            orphans.append(redline)
            continue
        if any(w in rewritten for w in w_indices):
            orphans.append(redline)
            continue
        rewritten.update(w_indices)

        _apply_diff_to_paragraph(
            targets[0], redline.original_text or "", redline.proposed_text or "", rev
        )
        for extra in targets[1:]:
            _delete_paragraph(extra, rev)

        if redline.rationale:
            _add_comment(
                targets[0], comments_root, next_cid, author, redline.rationale
            )
            next_cid += 1

    if orphans:
        next_cid = _append_orphans(body, comments_root, orphans, rev, author, next_cid)

    return _rezip(
        doc_root, comments_root, ct_bytes, rels_bytes, other_parts
    )


def _append_orphans(
    body: ET.Element,
    comments_root: ET.Element,
    orphans: list,
    rev: _Revision,
    author: str,
    next_cid: int,
) -> int:
    """Append clauses that have nowhere to sit in the original document.

    Covers MISSING findings - a protection the contract omits entirely has no
    paragraph to anchor to, but is often the most important thing in the review,
    so it must not be quietly dropped.
    """
    sect = body.find(f"{W}sectPr")
    if sect is not None:
        body.remove(sect)

    heading = ET.SubElement(body, f"{W}p")
    hppr = ET.SubElement(heading, f"{W}pPr")
    ET.SubElement(hppr, f"{W}pStyle").set(f"{W}val", "Heading1")
    wrapper = rev.ins()
    wrapper.append(_run("PROPOSED ADDITIONAL PROVISIONS", None))
    heading.append(wrapper)

    for redline in orphans:
        para = ET.SubElement(body, f"{W}p")
        label = redline.clause_title or redline.clause_type or "Additional provision"
        ins = rev.ins()
        ins.append(_run(f"{label}. {redline.proposed_text or ''}", None))
        para.append(ins)

        if redline.rationale:
            _add_comment(para, comments_root, next_cid, author, redline.rationale)
            next_cid += 1

    if sect is not None:
        body.append(sect)
    return next_cid


def _rezip(
    doc_root: ET.Element,
    comments_root: ET.Element,
    ct_bytes: bytes,
    rels_bytes: bytes,
    other_parts: dict,
) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zout:
        zout.writestr("[Content_Types].xml", ct_bytes)
        zout.writestr("word/_rels/document.xml.rels", rels_bytes)
        zout.writestr(
            "word/document.xml",
            ET.tostring(doc_root, xml_declaration=True, encoding="UTF-8"),
        )
        zout.writestr(
            "word/comments.xml",
            ET.tostring(comments_root, xml_declaration=True, encoding="UTF-8"),
        )
        for name, data in other_parts.items():
            zout.writestr(name, data)
    buffer.seek(0)
    return buffer.read()


def _lxml_run(text: str, deleted: bool = False):
    """python-docx builds on lxml, so the reconstructed path needs its own
    element factory - lxml will not adopt stdlib ElementTree nodes."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _lxml_revision(tag: str, rev: "_Revision"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    e = OxmlElement(f"w:{tag}")
    e.set(qn("w:id"), str(rev.take()))
    e.set(qn("w:author"), rev.author)
    e.set(qn("w:date"), rev.date)
    return e


def _apply_diff_lxml(para, original: str, proposed: str, rev: "_Revision") -> None:
    """Same interleaved diff as the .docx path, built with lxml elements."""
    p = para._p
    for child in list(p):
        if child.tag != f"{W}pPr":
            p.remove(child)

    for op in diff_ops(original, proposed):
        if op["op"] == "equal":
            p.append(_lxml_run(op["text"]))
        elif op["op"] == "delete":
            wrapper = _lxml_revision("del", rev)
            wrapper.append(_lxml_run(op["text"], deleted=True))
            p.append(wrapper)
        elif op["op"] == "insert":
            wrapper = _lxml_revision("ins", rev)
            wrapper.append(_lxml_run(op["text"]))
            p.append(wrapper)


def _lxml_comment_anchor(para, cid: int) -> None:
    """Mark a generated paragraph as the range of comment `cid`."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = para._p
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(cid))
    p.insert(0, start)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(cid))
    p.append(end)

    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(cid))
    ref_run.append(ref)
    p.append(ref_run)


def _inject_comments(docx_bytes: bytes, comments: list[tuple[int, str]], author: str) -> bytes:
    """Add a comments part to a python-docx file after the fact.

    python-docx cannot author comments, so the part is grafted on: the XML, the
    content-type override and the relationship all have to be present or Word
    drops every comment without warning.
    """
    if not comments:
        return docx_bytes

    root = ET.Element(f"{W}comments")
    for cid, text in comments:
        comment = _el("comment", id=cid, author=author, date=_now(), initials="AI")
        cpara = ET.SubElement(comment, f"{W}p")
        crun = ET.SubElement(cpara, f"{W}r")
        ctext = ET.SubElement(crun, f"{W}t")
        ctext.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        ctext.text = text
        root.append(comment)

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
        ct_bytes, rels_bytes = _register_comments_part(zin)
        parts = {
            name: zin.read(name)
            for name in zin.namelist()
            if name not in ("[Content_Types].xml", "word/_rels/document.xml.rels")
        }

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zout:
        zout.writestr("[Content_Types].xml", ct_bytes)
        zout.writestr("word/_rels/document.xml.rels", rels_bytes)
        zout.writestr(
            "word/comments.xml",
            ET.tostring(root, xml_declaration=True, encoding="UTF-8"),
        )
        for name, data in parts.items():
            zout.writestr(name, data)
    buffer.seek(0)
    return buffer.read()


def export_reconstructed_docx(review, version, redlines, author: str = "Contracts.AI") -> bytes:
    """Build a Word file from extracted blocks, for contracts uploaded as PDF.

    The text and the revision marks are faithful; the vendor's original layout
    is not recoverable. Callers surface that distinction in the UI.
    """
    from docx import Document

    blocks = load_blocks(version.blocks_json)
    exportable = _exportable(redlines)
    edits = {
        r.block_start: r for r in exportable if r.block_start is not None
    }
    superseded = set()
    for r in exportable:
        if r.block_start is None:
            continue
        for i in range(r.block_start + 1, (r.block_end or r.block_start) + 1):
            superseded.add(i)

    doc = Document()
    doc.add_heading(review.name or "Contract Redline", 0)
    note = doc.add_paragraph()
    note.add_run(
        "Reconstructed from a PDF upload. Tracked changes below are complete and "
        "accurate, but the original document's formatting could not be preserved."
    ).italic = True

    rev = _Revision(author)
    current_section = None
    comments: list[tuple[int, str]] = []
    next_cid = 1

    for block in blocks:
        if block["index"] in superseded:
            continue
        if block["section"] != current_section:
            current_section = block["section"]
            doc.add_heading(current_section, level=1)

        redline = edits.get(block["index"])
        para = doc.add_paragraph()

        if redline is None:
            para.add_run(block["text"])
            continue

        _apply_diff_lxml(
            para, redline.original_text or "", redline.proposed_text or "", rev
        )
        if redline.rationale:
            _lxml_comment_anchor(para, next_cid)
            comments.append((next_cid, redline.rationale))
            next_cid += 1

    orphans = [r for r in exportable if r.block_start is None]
    if orphans:
        doc.add_heading("Proposed Additional Provisions", level=1)
        for redline in orphans:
            para = doc.add_paragraph()
            label = redline.clause_title or redline.clause_type or "Additional provision"
            ins = _lxml_revision("ins", rev)
            ins.append(_lxml_run(f"{label}. {redline.proposed_text or ''}"))
            para._p.append(ins)
            if redline.rationale:
                _lxml_comment_anchor(para, next_cid)
                comments.append((next_cid, redline.rationale))
                next_cid += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return _inject_comments(buffer.read(), comments, author)


def export_redline_docx(
    review, version, redlines, author: str = "Contracts.AI"
) -> tuple[bytes, bool]:
    """Return (file bytes, is_faithful). `is_faithful` is False for PDF sources."""
    if version.doc_kind == "docx":
        return export_tracked_docx(version, redlines, author), True
    return export_reconstructed_docx(review, version, redlines, author), False
