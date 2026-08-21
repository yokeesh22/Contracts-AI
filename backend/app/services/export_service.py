import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CLASSIFICATION_COLORS = {
    "COMPLIANT": RGBColor(0x00, 0x80, 0x00),           # Green
    "ACCEPTABLE_DEVIATION": RGBColor(0xFF, 0x8C, 0x00), # Orange
    "CRITICAL_DEVIATION": RGBColor(0xCC, 0x00, 0x00),   # Red
    "NOT_APPLICABLE": RGBColor(0x80, 0x80, 0x80),        # Gray
}

CLASSIFICATION_LABELS = {
    "COMPLIANT": "Compliant",
    "ACCEPTABLE_DEVIATION": "Acceptable Deviation",
    "CRITICAL_DEVIATION": "Critical Deviation",
    "NOT_APPLICABLE": "Not Applicable",
}


def _set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_cell(cell, text: str, font_size: int = 9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)


def generate_exception_list_docx(session, requirements: list) -> bytes:
    """Generate a Customer Exception List Word document and return bytes."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title
    title = doc.add_heading("Customer Exception List", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sub-header info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Document Reference", session.urs_name),
        ("Specification", session.specification.name),
        ("Analysis Date", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Total Requirements", str(session.total_requirements)),
    ]
    for i, (label, value) in enumerate(info_data):
        _bold_cell(info_table.rows[i].cells[0], label)
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Summary counts
    counts = {k: 0 for k in CLASSIFICATION_LABELS}
    for req in requirements:
        cls = req.classification or "NOT_APPLICABLE"
        if cls in counts:
            counts[cls] += 1

    summary_heading = doc.add_heading("Summary", level=2)
    summary_table = doc.add_table(rows=1, cols=len(CLASSIFICATION_LABELS))
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0]
    bg_map = {
        "COMPLIANT": "00CC00",
        "ACCEPTABLE_DEVIATION": "FF8C00",
        "CRITICAL_DEVIATION": "CC0000",
        "NOT_APPLICABLE": "808080",
    }
    for i, (cls, label) in enumerate(CLASSIFICATION_LABELS.items()):
        cell = hdr.cells[i]
        _set_cell_bg(cell, bg_map[cls])
        cell.text = f"{label}\n{counts[cls]}"
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    doc.add_paragraph()

    # Main exception table
    doc.add_heading("Requirement Exception Details", level=2)
    headers = [
        "#",
        "Req. No.",
        "Customer Requirement",
        "Classification",
        "Specification Reference",
        "Deviation / Exception Detail",
        "Remarks",
    ]
    col_widths = [Cm(0.8), Cm(1.5), Cm(5), Cm(2.5), Cm(3), Cm(4.5), Cm(3)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "1F4E79")
        _bold_cell(cell, h, font_size=9)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.width = w

    # Data rows
    for idx, req in enumerate(requirements, start=1):
        cls = req.classification or "NOT_APPLICABLE"
        row = table.add_row()

        values = [
            str(idx),
            req.req_number or "-",
            req.req_text,
            CLASSIFICATION_LABELS.get(cls, cls),
            req.spec_reference or "-",
            req.deviation_detail or "-",
            req.remarks or "-",
        ]

        for i, (val, w) in enumerate(zip(values, col_widths)):
            cell = row.cells[i]
            cell.text = val
            cell.width = w
            para = cell.paragraphs[0]
            para.runs[0].font.size = Pt(8)

            # Shade classification cell
            if i == 3 and cls in bg_map:
                light_bg = {
                    "COMPLIANT": "E2EFDA",
                    "ACCEPTABLE_DEVIATION": "FCE4D6",
                    "CRITICAL_DEVIATION": "FFDCE1",
                    "NOT_APPLICABLE": "EDEDED",
                }
                _set_cell_bg(cell, light_bg.get(cls, "FFFFFF"))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
